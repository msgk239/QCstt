from fastapi import HTTPException
import os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from starlette.responses import FileResponse
from ..logger import get_logger
from .service import file_service
import json
import re

logger = get_logger(__name__)

class ExportService:
    """导出服务类，处理不同格式的文件导出"""
    
    def __init__(self):
        self.supported_formats = {
            'word': self._export_to_word,
            'pdf': self._export_to_pdf,
            'txt': self._export_to_txt,
            'md': self._export_to_markdown,
            'srt': self._export_to_srt,
        }
        # 创建导出文件的临时目录，设置为项目根目录下的exports
        self.export_dir = os.path.join(os.path.dirname(__file__), '..', '..', '..', 'exports')
        logger.info(f"导出目录设置为: {self.export_dir}")
        os.makedirs(self.export_dir, exist_ok=True)
    
    async def handle_export_request(self, file_id: str, format: str) -> FileResponse:
        """
        处理导出请求
        Args:
            file_id: 文件ID
            format: 导出格式
        Returns:
            FileResponse: 文件响应
        """
        try:
            logger.info(f"开始导出文件 {file_id} 为 {format} 格式")
            
            # 获取转写数据
            transcript = file_service.get_recognition_result(file_id)
            #logger.debug(f"获取到的转写数据: {transcript}")
            
            if transcript.get('code') != 200:
                raise HTTPException(status_code=400, detail="获取转写数据失败")
                
            # 使用导出服务处理导出
            file_path = self.export_transcript(file_id, format, transcript)
            logger.info(f"导出成功，文件路径: {file_path}")
            
            # 返回文件响应
            return FileResponse(
                path=file_path,
                media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                filename=f"transcript.{format}",
                background=None
            )
                
        except Exception as e:
            logger.error(f"导出失败: {str(e)}", exc_info=True)
            raise HTTPException(status_code=500, detail=str(e))
    
    def export_transcript(self, file_id: str, format: str, transcript_data: dict) -> str:
        """
        导出转写内容为指定格式
        Args:
            file_id: 文件ID
            format: 导出格式
            transcript_data: 转写数据
        Returns:
            str: 导出文件的路径
        """
        try:
            if format not in self.supported_formats:
                raise ValueError(f"不支持的导出格式: {format}")
                
            # 调用对应格式的导出方法
            export_func = self.supported_formats[format]
            return export_func(file_id, transcript_data)
            
        except Exception as e:
            logger.error(f"导出失败: {str(e)}", exc_info=True)
            raise HTTPException(status_code=500, detail=str(e))
    
    def _export_to_word(self, file_id: str, data: dict) -> str:
        """
        导出为Word文档
        Args:
            file_id: 文件ID
            data: 转写数据，格式如 original.json
        Returns:
            str: 导出文件的路径
        """
        try:
            # 创建新的 Word 文档
            doc = Document()
            
            # 从根目录的 metadata.json 获取文件名
            metadata_path = os.path.join('storage', 'metadata.json')
            logger.debug(f"读取 metadata 文件: {metadata_path}")  # 改为debug级别
            
            with open(metadata_path, 'r', encoding='utf-8') as f:
                metadata = json.load(f)
                #logger.debug(f"读取到的 metadata: {metadata}")  # 改为debug级别
            
            # 查找匹配的文件信息
            display_name = '转写文本'  # 默认标题
            for key in metadata:
                #logger.debug(f"正在检查键: {key}")  # 改为debug级别
                if key.startswith(file_id):
                    file_info = metadata[key]
                    display_name = file_info['display_name']  # 直接获取 display_name
                    logger.debug(f"找到匹配的文件名: {display_name}")  # 改为debug级别
                    break
            
            # 设置文档标题
            logger.debug(f"设置文档标题: {display_name}")  # 改为debug级别
            title = doc.add_heading(display_name, level=1)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 从 data 中获取实际的转写数据
            transcript_data = data.get('data', {}).get('data', {})
            segments = transcript_data.get('segments', [])
            
            # 合并连续的相同说话人的段落
            merged_segments = []
            current_segment = None
            
            for segment in segments:
                if not current_segment:
                    current_segment = {
                        'speakerDisplayName': segment['speakerDisplayName'],
                        'start_time': segment['start_time'],
                        'end_time': segment['end_time'],
                        'texts': [segment['text']]
                    }
                elif current_segment['speakerDisplayName'] == segment['speakerDisplayName']:
                    # 如果说话人相同，合并文本并更新结束时间
                    current_segment['texts'].append(segment['text'])
                    current_segment['end_time'] = segment['end_time']
                else:
                    # 说话人不同，保存当前段落并开始新段落
                    merged_segments.append(current_segment)
                    current_segment = {
                        'speakerDisplayName': segment['speakerDisplayName'],
                        'start_time': segment['start_time'],
                        'end_time': segment['end_time'],
                        'texts': [segment['text']]
                    }
            
            # 添加最后一个段落
            if current_segment:
                merged_segments.append(current_segment)
            
            # 将合并后的段落写入文档
            for segment in merged_segments:
                # 添加段落
                p = doc.add_paragraph()
                
                # 添加说话人名称（加粗）
                speaker_run = p.add_run(f"{segment['speakerDisplayName']}: ")
                speaker_run.bold = True
                
                # 添加时间戳（灰色小字）
                start_time = self._format_time(segment['start_time'])
                end_time = self._format_time(segment['end_time'])
                time_run = p.add_run(f"[{start_time} - {end_time}]")
                time_run.font.size = Pt(9)
                time_run.font.color.rgb = RGBColor(128, 128, 128)
                
                # 添加合并后的文本内容
                text_content = '\n'.join(segment['texts'])
                p.add_run(f"\n{text_content}\n")
            
            # 生成导出文件路径
            export_path = os.path.join(self.export_dir, f"{file_id}_transcript.docx")
            
            # 添加日志查看文档内容
            logger.debug("准备保存文档内容")  # 改为debug级别
            #for paragraph in doc.paragraphs:
                #logger.debug(f"段落内容: {paragraph.text}")  # 改为debug级别
            
            # 保存文档
            doc.save(export_path)
            logger.info(f"Word文档导出成功: {export_path}")  # 保留为info级别，这是关键信息
            
            return export_path
            
        except Exception as e:
            logger.error(f"Word导出失败: {str(e)}", exc_info=True)
            raise Exception(f"Word导出失败: {str(e)}")
    
    def _format_time(self, seconds: float) -> str:
        """
        将秒数转换为时间字符串格式 (mm:ss)
        """
        minutes = int(seconds // 60)
        seconds = int(seconds % 60)
        return f"{minutes:02d}:{seconds:02d}"
    
    def _export_to_pdf(self, file_id: str, data: dict) -> str:
        """导出为PDF文件"""
        # TODO: 实现PDF导出逻辑
        raise NotImplementedError("PDF导出功能尚未实现")
    
    def _export_to_txt(self, file_id: str, data: dict) -> str:
        """导出为纯文本文件"""
        # TODO: 实现TXT导出逻辑
        raise NotImplementedError("TXT导出功能尚未实现")
    
    def _export_to_markdown(self, file_id: str, data: dict) -> str:
        """导出为Markdown文件"""
        # TODO: 实现Markdown导出逻辑
        raise NotImplementedError("Markdown导出功能尚未实现")
    
    def _export_to_srt(self, file_id: str, data: dict) -> str:
        """
        导出为SRT字幕文件
        Args:
            file_id: 文件ID
            data: 转写数据，格式如 original.json
        Returns:
            str: 导出文件的路径
        """
        try:
            # 从 data 中获取实际的转写数据
            transcript_data = data.get('data', {}).get('data', {})
            segments = transcript_data.get('segments', [])

            # 生成导出文件路径，使用 .srt 扩展名
            export_path = os.path.join(self.export_dir, f"{file_id}_transcript.srt")

            with open(export_path, 'w', encoding='utf-8') as srt_file:
                index = 1

                for segment in segments:
                    # 检查是否有子段落，前端可能会传入合并后的段落
                    if 'subSegments' in segment and segment['subSegments']:
                        # 使用子段落
                        current_index = index
                        for sub_segment in segment['subSegments']:
                           processed_lines_count = self._process_segment_to_srt(sub_segment, srt_file, current_index)
                           current_index += processed_lines_count # 累加处理的行数
                        index = current_index # 更新主索引
                    else:
                        # 使用主段落
                        processed_lines_count = self._process_segment_to_srt(segment, srt_file, index)
                        index += processed_lines_count # 更新索引

            logger.info(f"SRT文档导出成功: {export_path}")
            return export_path

        except Exception as e:
            logger.error(f"SRT导出失败: {str(e)}", exc_info=True)
            raise Exception(f"SRT导出失败: {str(e)}")

    def _remove_punctuation(self, text: str) -> str:
        """移除文本中的标点符号（。，）"""
        return re.sub(r'[。，]', '', text)

    def _process_segment_to_srt(self, segment, srt_file, start_index) -> int:
        """
        处理一个段落并转换为SRT格式
        Returns:
            int: 处理并写入的SRT条目数量
        """
        processed_lines_count = 0
        # 检查是否有字级别的时间戳
        if 'timestamps' in segment and segment['timestamps'] and len(segment['timestamps']) > 0:
            # 处理字级别的时间戳
            processed_lines_count = self._process_with_char_timestamps(segment, srt_file, start_index)
        else:
            # 无字级别时间戳，使用段落级别的时间戳
            processed_lines_count = self._process_without_char_timestamps(segment, srt_file, start_index)
        return processed_lines_count

    def _process_with_char_timestamps(self, segment, srt_file, start_index):
        """
        使用字级别时间戳处理段落
        Returns:
            int: 处理并写入的SRT条目数量
        """
        timestamps = segment.get('timestamps', [])
        text = segment.get('text', '')
        max_length = 30 # 定义最大长度，与 _split_text 保持一致或根据需要调整

        if not timestamps or not text:
            return 0 # 如果没有时间戳或文本，无法处理

        # 1. 初步根据标点分割，并保留时间戳
        initial_parts = []
        current_part_text = ""
        current_part_timestamps = []

        for i, char in enumerate(text):
            current_part_text += char
            # 确保索引不越界
            if i < len(timestamps):
                current_part_timestamps.append(timestamps[i])
            else:
                # 如果时间戳数量少于字符数，记录警告并可能需要跳过后续处理
                logger.warning(f"时间戳数量 ({len(timestamps)}) 少于字符数 ({len(text)}) in segment, potential issue.")
                # 可以选择补充默认时间戳或中断处理，这里暂时继续，但可能导致后续错误
                if current_part_timestamps:
                     last_ts = current_part_timestamps[-1]
                     current_part_timestamps.append({'start': last_ts['end'], 'end': last_ts['end'] + 0.1}) # 估算一个时间
                else:
                     # 如果一开始就没有时间戳，无法处理
                     return 0


            # 根据标点分割
            if char in '。，':
                if current_part_text.strip() and current_part_timestamps:
                    initial_parts.append({
                        'text': current_part_text.strip(),
                        'timestamps': list(current_part_timestamps), # 复制列表
                        'start_time': current_part_timestamps[0]['start'],
                        'end_time': current_part_timestamps[-1]['end']
                    })
                current_part_text = ""
                current_part_timestamps.clear()

        # 添加最后剩余的部分
        if current_part_text.strip() and current_part_timestamps:
             initial_parts.append({
                 'text': current_part_text.strip(),
                 'timestamps': list(current_part_timestamps),
                 'start_time': current_part_timestamps[0]['start'],
                 'end_time': current_part_timestamps[-1]['end']
             })

        # 2. 处理初步分割结果，对超长部分进行智能分割
        final_parts = []
        for part in initial_parts:
            if len(part['text']) <= max_length:
                # 时间戳信息不再需要，只保留起止时间
                final_parts.append({
                    'text': part['text'],
                    'start_time': part['start_time'],
                    'end_time': part['end_time']
                })
            else:
                # 调用新方法进行智能分割
                split_sub_parts = self._split_long_part_by_timestamp(part, max_length)
                final_parts.extend(split_sub_parts)

        # 3. 写入SRT文件
        index = start_index
        processed_count = 0
        for part in final_parts:
            srt_start = self._format_srt_time(part['start_time'])
            srt_end = self._format_srt_time(part['end_time'])
            # 移除标点符号后再写入
            clean_text = self._remove_punctuation(part['text'])
            if clean_text: # 确保移除标点后仍有内容
                srt_file.write(f"{index}\n")
                srt_file.write(f"{srt_start} --> {srt_end}\n")
                srt_file.write(f"{clean_text}\n\n")
                index += 1
                processed_count += 1
        return processed_count

    def _is_safe_split_boundary(self, char_before: str, char_after: str) -> bool:
        """
        检查在两个字符之间分割是否安全（避免打断英文单词或数字）
        如果两个字符都是ASCII字母或数字，则认为不安全。
        """
        # 简单的检查，可以根据需要扩展（例如考虑连字符等）
        char_before_is_alnum = char_before.isalnum() and char_before.isascii()
        char_after_is_alnum = char_after.isalnum() and char_after.isascii()

        # 如果前后都是字母/数字，则不安全
        if char_before_is_alnum and char_after_is_alnum:
            return False
        return True

    def _split_long_part_by_timestamp(self, part: dict, max_length: int) -> list:
        """
        将包含时间戳的长文本部分根据最大时间间隔智能分割，并考虑词语边界。
        Args:
            part: 包含 'text', 'timestamps', 'start_time', 'end_time' 的字典
            max_length: 每部分的最大长度
        Returns:
            list: 分割后的部分列表，每个元素是 {'text': str, 'start_time': float, 'end_time': float}
        """
        result_parts = []
        current_text = part['text']
        current_timestamps = part['timestamps']

        while len(current_text) > max_length:
            best_safe_split_index = -1
            max_safe_gap = -1.0
            best_unsafe_split_index = -1 # 用于回退
            max_unsafe_gap = -1.0     # 用于回退

            search_end_index = min(max_length, len(current_timestamps))
            if search_end_index <= 10:
                split_index = max_length # 范围太小，直接硬分割
            else:
                for i in range(10, search_end_index):
                    if i < len(current_timestamps) and i > 0 and i < len(current_text): # 确保索引有效
                         try:
                             gap = current_timestamps[i]['start'] - current_timestamps[i-1]['end']
                             char_before = current_text[i-1]
                             char_after = current_text[i]

                             is_safe = self._is_safe_split_boundary(char_before, char_after)

                             if is_safe:
                                 if gap > max_safe_gap:
                                     max_safe_gap = gap
                                     best_safe_split_index = i
                             else: # 不安全的分割点也记录下来，用于回退
                                 if gap > max_unsafe_gap:
                                     max_unsafe_gap = gap
                                     best_unsafe_split_index = i

                         except (KeyError, TypeError, IndexError) as e:
                             logger.warning(f"计算或检查时间戳/文本边界时出错 at index {i}: {e}.")
                             continue
                    else:
                        break # 索引无效

                # 决定最终分割点
                if best_safe_split_index != -1: # 优先使用安全且间隔最大的点
                    split_index = best_safe_split_index
                elif best_unsafe_split_index != -1: # 回退：使用不安全但间隔最大的点
                    split_index = best_unsafe_split_index
                    logger.debug(f"未找到安全分割点, 回退使用最大间隔点 (不安全): index {split_index}")
                else: # 最后回退：硬分割
                    split_index = max_length
                    logger.debug(f"未找到任何合适间隔点, 回退使用 max_length: index {split_index}")


            # --- 后续分割逻辑保持不变 ---
            # 确保 split_index 不超过当前文本长度
            split_index = min(split_index, len(current_text))
            # 确保分割后至少有一个字符
            if split_index == 0: split_index = 1

            # 确保分割索引对应的时间戳存在
            if split_index > 0 and split_index - 1 < len(current_timestamps) and len(current_timestamps) > 0:
                 new_part_text = current_text[:split_index]
                 # 确保使用正确的时间戳索引
                 if split_index > 0 and split_index - 1 < len(current_timestamps):
                     new_part_start_time = current_timestamps[0]['start']
                     new_part_end_time = current_timestamps[split_index-1]['end']

                     result_parts.append({
                         'text': new_part_text,
                         'start_time': new_part_start_time,
                         'end_time': new_part_end_time
                     })

                     # 更新剩余部分
                     current_text = current_text[split_index:]
                     current_timestamps = current_timestamps[split_index:]

                     # 如果时间戳耗尽，处理剩余文本
                     if current_text and not current_timestamps:
                         logger.warning("文本分割后时间戳耗尽，剩余文本将无法处理。")
                         last_known_end_time = result_parts[-1]['end_time'] if result_parts else part['start_time']
                         estimated_end_time = last_known_end_time + (len(current_text) * 0.1)
                         result_parts.append({
                             'text': current_text,
                             'start_time': last_known_end_time,
                             'end_time': estimated_end_time
                         })
                         current_text = "" # 标记为已处理
                         break
                 else:
                     # 如果split_index-1无效，这是个错误情况
                      logger.error(f"分割逻辑错误：split_index {split_index} 导致时间戳索引越界。")
                      # 处理错误，例如将整个剩余部分作为一个块
                      if current_text and current_timestamps:
                          result_parts.append({
                              'text': current_text,
                              'start_time': current_timestamps[0]['start'],
                              'end_time': current_timestamps[-1]['end']
                          })
                      elif current_text:
                          last_known_end_time = result_parts[-1]['end_time'] if result_parts else part['start_time']
                          estimated_end_time = last_known_end_time + (len(current_text) * 0.1)
                          result_parts.append({
                                'text': current_text,
                                'start_time': last_known_end_time,
                                'end_time': estimated_end_time
                          })
                      current_text = "" # 标记为已处理
                      break


            else:
                # 如果时间戳索引无效，可能是数据问题或逻辑错误
                logger.error(f"无法根据 split_index={split_index} 分割，时间戳或文本长度不足。")
                if current_text and current_timestamps:
                     result_parts.append({
                         'text': current_text,
                         'start_time': current_timestamps[0]['start'],
                         'end_time': current_timestamps[-1]['end']
                     })
                elif current_text:
                    last_known_end_time = result_parts[-1]['end_time'] if result_parts else part['start_time']
                    estimated_end_time = last_known_end_time + (len(current_text) * 0.1)
                    result_parts.append({
                          'text': current_text,
                          'start_time': last_known_end_time,
                          'end_time': estimated_end_time
                    })
                current_text = ""
                break


        # 添加最后剩余的部分
        if current_text and current_timestamps:
            result_parts.append({
                'text': current_text,
                'start_time': current_timestamps[0]['start'],
                'end_time': current_timestamps[-1]['end']
            })
        elif current_text: # 处理上面break后可能未处理的剩余文本
             if not any(p['text'] == current_text for p in result_parts): # 避免重复添加
                last_known_end_time = result_parts[-1]['end_time'] if result_parts else part['start_time']
                estimated_end_time = last_known_end_time + (len(current_text) * 0.1)
                result_parts.append({
                    'text': current_text,
                    'start_time': last_known_end_time,
                    'end_time': estimated_end_time
                })


        return result_parts

    def _process_without_char_timestamps(self, segment, srt_file, start_index):
        """
        使用段落级别时间戳处理段落
        Returns:
            int: 处理并写入的SRT条目数量
        """
        start_time = self._format_srt_time(segment['start_time'])
        end_time = self._format_srt_time(segment['end_time'])
        text_content = segment['text']
        
        # 分割文本为多行
        lines = self._split_text(text_content)
        index = start_index
        processed_count = 0

        for line in lines:
            clean_line = line.strip()
            if clean_line:  # 确保不写入空行
                # 移除标点符号后再写入
                clean_text = self._remove_punctuation(clean_line)
                if clean_text: # 确保移除标点后仍有内容
                    srt_file.write(f"{index}\n")
                    srt_file.write(f"{start_time} --> {end_time}\n")
                    srt_file.write(f"{clean_text}\n\n")
                    index += 1
                    processed_count += 1
        return processed_count
    
    def _format_srt_time(self, seconds: float) -> str:
        """
        将秒数转换为SRT时间字符串格式 (hh:mm:ss,ms)
        确保毫秒部分的精度准确
        """
        if seconds is None:
            seconds = 0.0
            
        hours = int(seconds // 3600)
        minutes = int((seconds % 3600) // 60)
        seconds_value = seconds % 60
        # 保留整数部分和小数部分，确保毫秒的精度
        seconds_int = int(seconds_value)
        milliseconds = int(round((seconds_value - seconds_int) * 1000))
        
        return f"{hours:02}:{minutes:02}:{seconds_int:02},{milliseconds:03}"
    
    def _split_text(self, text: str, max_length: int = 30) -> list:
        """
        分割文本为多行，在逗号和句号处分割，同时保留标点符号，并限制每行的长度
        """
        # 使用正则表达式在标点符号后进行分割，但保留标点符号
        pattern = r'([。，])'
        parts = re.split(pattern, text)
        
        lines = []
        current_line = ""
        
        i = 0
        while i < len(parts):
            part = parts[i]
            
            # 如果当前部分是标点符号
            if i + 1 < len(parts) and parts[i + 1] in '。，':
                punctuation = parts[i + 1]
                combined = part + punctuation
                
                # 如果添加这部分后长度超过限制，先添加当前行
                if current_line and len(current_line) + len(combined) > max_length:
                    lines.append(current_line)
                    current_line = combined
                else:
                    current_line += combined
                
                i += 2  # 跳过已处理的标点符号
            else:
                # 普通文本部分
                if current_line and len(current_line) + len(part) > max_length:
                    lines.append(current_line)
                    current_line = part
                else:
                    current_line += part
                i += 1
            
            # 如果当前部分以标点符号结尾，将其作为一个独立的行
            if current_line and current_line[-1] in '。，':
                lines.append(current_line)
                current_line = ""
        
        # 添加最后剩余的行
        if current_line:
            lines.append(current_line)
            
        return lines

# 创建全局实例
export_service = ExportService() 